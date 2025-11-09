# docx_handler.py
from docx import Document
from docx.shared import Inches
import io

class DocxGenerator:
    def __init__(self):
        self.document = Document()
    
    def generate_from_template(self, template_text, variables):
        """Генерация DOCX документа из шаблона"""
        doc = Document()
        
        # Замена переменных в тексте
        formatted_text = template_text
        for key, value in variables.items():
            formatted_text = formatted_text.replace(f"{{{key}}}", str(value))
        
        # Добавление текста в документ
        for paragraph in formatted_text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        return doc
    
    def save_to_buffer(self, doc):
        """Сохранение документа в буфер"""
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer